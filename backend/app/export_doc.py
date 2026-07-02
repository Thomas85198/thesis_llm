"""Editor-mode export: TipTap (ProseMirror) JSON → .docx / .tex.

Walks the stored document tree and renders block/inline nodes, including the
custom inline `citation` atom, into Word and LaTeX. In-text citation markers and
the trailing reference list mirror frontend/lib/citation-format.ts so an export
reads exactly like what the author sees on screen.
"""

from __future__ import annotations

import base64
import io
import ipaddress
import os
import re
import socket
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------- image fetching (for embedding figures in export) ----------

_IMG_EXTS = {"png", "jpg", "jpeg", "gif", "webp"}
_IMG_PATH_RE = re.compile(r"/api/editor/images/([\w.\-]+)$")
_MAX_IMG_BYTES = 10 * 1024 * 1024


def _upload_dir() -> Path:
    # Mirror routes._resolve_upload_dir (avoid a circular import).
    explicit = os.getenv("UPLOAD_DIR")
    if explicit:
        return Path(explicit)
    data_dir = os.getenv("DATA_DIR")
    if data_dir:
        return Path(data_dir) / "uploads"
    return Path(__file__).parent.parent / "uploads"


def _is_public_http_url(src: str) -> bool:
    """SSRF guard for figure fetches: user-controlled srcs reach a server-side
    GET, which must not be able to read compose-internal services
    (http://neo4j:7474, cloud metadata, …)."""
    try:
        parsed = urlparse(src)
        if parsed.scheme not in ("http", "https") or not parsed.hostname:
            return False
        for info in socket.getaddrinfo(parsed.hostname, None):
            ip = ipaddress.ip_address(info[4][0])
            if not ip.is_global:
                return False
        return True
    except Exception:  # noqa: BLE001 — unresolvable host etc. → treat as unsafe
        return False


def _image_bytes(src: str) -> tuple[bytes, str] | None:
    """Resolve a figure src to (bytes, ext). Reads our own uploads from disk;
    fetches public external image URLs over HTTP. None on any failure."""
    if not src:
        return None
    try:
        m = _IMG_PATH_RE.search(src)
        if m:
            p = _upload_dir() / m.group(1)
            if not p.is_file():
                return None
            ext = p.suffix.lstrip(".").lower() or "png"
            return p.read_bytes(), (ext if ext in _IMG_EXTS else "png")
        if not _is_public_http_url(src):
            return None
        # No redirects (a public URL could 302 to an internal one) and a
        # streamed size cap (the old post-download check still buffered the
        # whole body in memory first).
        with httpx.Client(timeout=15.0, follow_redirects=False) as cli:
            with cli.stream("GET", src) as resp:
                resp.raise_for_status()
                chunks: list[bytes] = []
                size = 0
                for chunk in resp.iter_bytes():
                    size += len(chunk)
                    if size > _MAX_IMG_BYTES:
                        return None
                    chunks.append(chunk)
                data = b"".join(chunks)
        ext = src.rsplit(".", 1)[-1].split("?")[0].lower() if "." in src else "png"
        return data, (ext if ext in _IMG_EXTS else "png")
    except Exception:  # noqa: BLE001 — best-effort; fall back to caption text
        return None


def _img_data_uri(src: str) -> str:
    """Inline an image as a base64 data URI so a standalone HTML export (online
    preview / print-to-PDF / downloaded .html) carries the picture itself — the
    raw /api/editor/images/ path can't resolve outside the app's own origin."""
    img = _image_bytes(src)
    if not img:
        return src
    data, ext = img
    mime = "jpeg" if ext == "jpg" else ext
    return f"data:image/{mime};base64,{base64.b64encode(data).decode('ascii')}"


# ---------- citation formatting (mirror frontend lib/citation-format.ts) ----------


def _author_list(authors: str) -> list[str]:
    return [a.strip() for a in (authors or "").split(",") if a.strip()]


def _last_name(full: str) -> str:
    parts = full.strip().split()
    return parts[-1] if parts else full


def is_numbered(style: str) -> bool:
    """Styles that render in-text as a bracketed number rather than author–year."""
    return style in ("ieee", "numeric")


def _authors_short(authors: str) -> str:
    lst = _author_list(authors)
    if not lst:
        return "Anon."
    if len(lst) == 1:
        return _last_name(lst[0])
    if len(lst) == 2:
        return f"{_last_name(lst[0])} & {_last_name(lst[1])}"
    return f"{_last_name(lst[0])} et al."


def in_text_label(attrs: dict, style: str, number: int) -> str:
    short = _authors_short(attrs.get("authors", ""))
    y = attrs.get("year") or "n.d."
    # Narrative: the author is named in the sentence, so only the year (or [n])
    # is parenthesized — "Vaswani et al. (2017)" / "Vaswani et al. [3]".
    if attrs.get("narrative"):
        if is_numbered(style):
            return f"{short} [{number}]"
        if style == "mla":
            return short
        return f"{short} ({y})"
    if is_numbered(style):
        return f"[{number}]"
    if style == "mla":
        return f"({short})"
    if style == "chicago":
        return f"({short} {y})"
    return f"({short}, {y})"  # apa, harvard


def full_reference(attrs: dict, style: str, number: int) -> str:
    authors = attrs.get("authors") or "Anon."
    year = attrs.get("year") or "n.d."
    title = attrs.get("title") or "(untitled)"
    venue = attrs.get("venue") or ""
    if style == "mla":
        return f"{authors}. “{title}.”" + (f" {venue}," if venue else "") + f" {year}."
    if style == "chicago":
        return f"{authors}. “{title}.”" + (f" {venue}" if venue else "") + f" ({year})."
    if style == "harvard":
        return f"{authors} ({year}) ‘{title}’" + (f", {venue}" if venue else "") + "."
    if style in ("ieee", "numeric"):
        return (
            f"[{number}] {authors}, “{title},”"
            + (f" {venue}," if venue else "")
            + f" {year}."
        )
    return f"{authors} ({year}). {title}" + (f". {venue}" if venue else "") + "."  # apa


# ---------- citation identity / grouping (Taiwan reference list) ----------

# Han (incl. Ext-A) + kana — used to route a reference to the 中文 / 英文 group.
_CJK_RE = re.compile(r"[぀-ヿ㐀-䶿一-鿿]")


def _cite_key(a: dict) -> str:
    """Stable identity mirroring the editor's citeKey(): the OpenAlex id when
    linked, else a normalized authors|year key — so Crossref/manual/unlinked
    citations number and dedup correctly (not just OpenAlex-backed ones)."""
    oid = a.get("openalexId")
    if oid:
        return oid
    authors = " ".join((a.get("authors") or "").lower().split()).strip()
    year = a.get("year")
    return f"{authors}|{year if year is not None else ''}"


def _ref_group(a: dict) -> str:
    """Route a reference into the Taiwan three-way split: 網路資源 / 中文 / 英文.
    A web resource = explicitly kind=web, or a bare URL with no DOI/venue."""
    if a.get("kind") == "web":
        return "web"
    if a.get("url") and not a.get("doi") and not a.get("venue"):
        return "web"
    blob = f"{a.get('title', '')}{a.get('authors', '')}{a.get('venue', '')}"
    return "zh" if _CJK_RE.search(blob) else "en"


def _ref_sort_key(a: dict) -> str:
    """Within-group ordering: by first author (surname-ish), case-folded."""
    return (a.get("authors") or "").strip().lower()


def _tw_reference(a: dict) -> str:
    """One Taiwan-thesis reference entry as plain text (DOCX / preview). Chinese
    entries use full-width punctuation （）。and 、 between authors; English
    entries follow APA. DOI / URL appended per group."""
    authors = a.get("authors") or "Anon."
    year = a.get("year") or "n.d."
    title = a.get("title") or "(untitled)"
    venue = a.get("venue") or ""
    url = a.get("url") or ""
    doi = a.get("doi") or ""
    is_cjk = bool(_CJK_RE.search(f"{authors}{title}{venue}"))
    if is_cjk:
        authors = authors.replace(", ", "、")
        s = f"{authors}（{year}）。{title}"
        if venue:
            s += f"。{venue}"
        s += "。"
        if doi:
            s += f"doi: {doi}"
        elif url:
            s += f"擷取自 {url}"
        return s
    s = f"{authors} ({year}). {title}"
    if venue:
        s += f". {venue}"
    s += "."
    if doi:
        s += f" doi: {doi}"
    elif url:
        s += f" Retrieved from {url}"
    return s


def _tw_reference_latex(a: dict) -> str:
    """LaTeX form of _tw_reference: the body is escaped, but DOI / URL go through
    \\url{} so a long link breaks across lines instead of overflowing the margin."""
    authors = a.get("authors") or "Anon."
    year = str(a.get("year") or "n.d.")
    title = a.get("title") or "(untitled)"
    venue = a.get("venue") or ""
    url = a.get("url") or ""
    doi = a.get("doi") or ""
    is_cjk = bool(_CJK_RE.search(f"{authors}{title}{venue}"))
    if is_cjk:
        authors = authors.replace(", ", "、")
        s = f"{_esc(authors)}（{_esc(year)}）。{_esc(title)}"
        if venue:
            s += f"。{_esc(venue)}"
        s += "。"
        if doi:
            s += f"doi: \\url{{{doi}}}"
        elif url:
            s += f"擷取自 \\url{{{url}}}"
        return s
    s = f"{_esc(authors)} ({_esc(year)}). {_esc(title)}"
    if venue:
        s += f". {_esc(venue)}"
    s += "."
    if doi:
        s += f" doi: \\url{{{doi}}}"
    elif url:
        s += f" Retrieved from \\url{{{url}}}"
    return s


# ---------- tree walking ----------


def _children(node: dict) -> list[dict]:
    return node.get("content") or []


def collect_citations(doc: dict) -> list[dict]:
    """Distinct citation attrs by first appearance — the reference-list order."""
    seen: set[str] = set()
    out: list[dict] = []

    def walk(node: dict) -> None:
        if node.get("type") == "citation":
            a = node.get("attrs") or {}
            k = _cite_key(a)
            if k and k != "|" and k not in seen:
                seen.add(k)
                out.append(a)
        for ch in _children(node):
            walk(ch)

    walk(doc)
    return out


def _citation_number(order: list[str], oid: str) -> int:
    """1-based first-appearance number (mirrors the editor's numeric labels)."""
    return order.index(oid) + 1 if oid in order else len(order) + 1


def _block_text(block: dict) -> str:
    """Flatten a block's inline text (used for code blocks)."""
    return "".join(
        n.get("text", "") for n in _children(block) if n.get("type") == "text"
    )


def _plain_text(node: dict) -> str:
    """Recursively gather visible text from a node (used for table cells)."""
    if node.get("type") == "text":
        return node.get("text", "")
    return "".join(_plain_text(ch) for ch in _children(node))


# Directory headings are skipped in a TOC — it shouldn't list itself.
_DIR_HEADING_TITLES = {
    "目錄",
    "目次",
    "圖目錄",
    "圖次",
    "表目錄",
    "表次",
    "table of contents",
    "contents",
    "list of figures",
    "list of tables",
}

# Level-1 headings that are NOT numbered chapters in a Taiwan thesis (參考文獻 /
# 附錄 / 誌謝 / 摘要 …). They render as an unnumbered, centered chapter and don't
# advance the 第N章 counter; the body's Arabic page numbering also starts only at
# the first *numbered* chapter, so front matter (摘要/目錄) stays Roman.
_UNNUMBERED_CHAPTERS = {
    "參考文獻",
    "参考文献",
    "附錄",
    "附录",
    "誌謝",
    "志謝",
    "致謝",
    "謝誌",
    "中文摘要",
    "英文摘要",
    "摘要",
    "abstract",
    "references",
    "bibliography",
    "appendix",
    "appendices",
    "acknowledgements",
    "acknowledgments",
}


def _is_unnumbered_chapter(text: str) -> bool:
    return text.strip().lower() in _UNNUMBERED_CHAPTERS


def _collect_headings(doc: dict) -> list[tuple[int, str]]:
    """(level, text) for every content heading — source for a static table of contents."""
    out: list[tuple[int, str]] = []

    def walk(node: dict) -> None:
        if node.get("type") == "heading":
            text = _plain_text(node).strip()
            if text.lower() not in _DIR_HEADING_TITLES:
                out.append(((node.get("attrs") or {}).get("level", 1), text))
        for ch in _children(node):
            walk(ch)

    walk(doc)
    return out


# A standalone heading with one of these texts is redundant when a
# `referenceList` node is present — the node renders its own 「參考文獻」 title,
# so we drop the manual heading to avoid a duplicate (one empty, one with the
# list). Keeps LaTeX/PDF/DOCX consistent.
_REFS_HEADING_TITLES = {
    "參考文獻",
    "参考文献",
    "參考書目",
    "參考資料",
    "references",
    "bibliography",
}


def _prepared_content(document: dict) -> dict:
    """Expand each live `tableOfContents` node into a static, indented heading list
    so exports carry a real table of contents (the editor-only node has no body).
    Also drops a redundant 「參考文獻」 heading when a referenceList node exists."""
    raw = document.get("content_json") or {}
    if not isinstance(raw, dict):
        return {"type": "doc", "content": []}
    headings = _collect_headings(raw)
    top = raw.get("content") or []
    has_ref_list = any(b.get("type") == "referenceList" for b in top)
    blocks: list[dict] = []
    for b in top:
        if (
            has_ref_list
            and b.get("type") == "heading"
            and _plain_text(b).strip().lower() in _REFS_HEADING_TITLES
        ):
            continue  # the referenceList node supplies the 參考文獻 heading
        if b.get("type") == "tableOfContents":
            for level, text in headings:
                if not text:
                    continue
                indent = "　" * (max(level, 1) - 1)  # full-width space per level
                blocks.append(
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": f"{indent}{text}"}],
                    }
                )
        else:
            blocks.append(b)
    return {"type": "doc", "content": blocks}


def _table_parts(block: dict) -> tuple[str, list[list[str]]]:
    """Split a tableBlock into (caption, rows-of-cell-text)."""
    caption = ""
    rows: list[list[str]] = []
    for child in _children(block):
        if child.get("type") == "tableCaption":
            caption = _plain_text(child).strip()
        elif child.get("type") == "table":
            for tr in _children(child):
                rows.append([_plain_text(cell).strip() for cell in _children(tr)])
    return caption, rows


# ---------- DOCX ----------


_CJK_DIGITS = "〇一二三四五六七八九"


def _cjk_num(n: int) -> str:
    """1..99 → 一/十/十一/二十… for 「第N章」 chapter headings."""
    if n <= 0:
        return str(n)
    if n < 10:
        return _CJK_DIGITS[n]
    if n < 20:
        return "十" + (_CJK_DIGITS[n % 10] if n % 10 else "")
    if n < 100:
        tens, ones = divmod(n, 10)
        return _CJK_DIGITS[tens] + "十" + (_CJK_DIGITS[ones] if ones else "")
    return str(n)


def _apply_run_font(run) -> None:
    """Pin Times New Roman (ascii/hAnsi) + 標楷體 (eastAsia) on a run so inline
    CJK stays in 標楷體 even when the run also carries Latin text. The size is
    inherited from the Normal style (14pt for twthesis)."""
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "標楷體")


def _style_heading_runs(paragraph) -> None:
    """Thesis headings are black 標楷體 (Word's default heading style is a blue
    sans) — recolor + repin the font, keep the style's size/weight."""
    for r in paragraph.runs:
        _apply_run_font(r)
        r.font.color.rgb = RGBColor(0, 0, 0)


def _add_inline_docx(
    paragraph, nodes: list[dict], style: str, order: list[str]
) -> None:
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            run = paragraph.add_run(n.get("text", ""))
            marks = {m.get("type") for m in (n.get("marks") or [])}
            run.bold = "bold" in marks
            run.italic = "italic" in marks
            if "strike" in marks:
                run.font.strike = True
            if "code" in marks:
                run.font.name = "Courier New"
            else:
                _apply_run_font(run)
        elif t == "citation":
            a = n.get("attrs") or {}
            num = _citation_number(order, _cite_key(a))
            _apply_run_font(paragraph.add_run(in_text_label(a, style, num)))
        elif t == "mathInline":
            paragraph.add_run(f"${(n.get('attrs') or {}).get('latex', '')}$")


def _docx_caption_num(ctx: dict, kind: str) -> str:
    """Chapter-scoped caption number: 「圖 2-1」/「表 2-1」 when chapters exist,
    else a flat 「圖 1」. kind is 'fig' or 'tab'."""
    ctx[kind] = ctx.get(kind, 0) + 1
    chap = ctx.get("chapter", 0)
    return f"{chap}-{ctx[kind]}" if chap > 0 else str(ctx[kind])


def _render_block_docx(
    docx: Document, block: dict, style: str, order: list[str], ctx: dict
) -> None:
    tw = ctx.get("tw", False)
    t = block.get("type")
    level = (block.get("attrs") or {}).get("level", 1) if t == "heading" else 0
    numbered_chapter = (
        t == "heading"
        and tw
        and level == 1
        and not _is_unnumbered_chapter(_plain_text(block))
    )
    opens_body = bool(
        numbered_chapter
        and ctx.get("roman_split")
        and not ctx.get("body_section_started")
    )
    # The cover's trailing page break: a chapter consumes it via its section
    # break; any other first block (front matter) gets an explicit page break.
    if ctx.get("pending_cover_break"):
        ctx["pending_cover_break"] = False
        if not opens_body:
            docx.add_page_break()
    if t == "heading":
        # First numbered chapter: open the Arabic body section (Roman→Arabic).
        if opens_body:
            sec = docx.add_section(WD_SECTION.NEW_PAGE)
            _docx_section_margins(sec)
            sec.different_first_page_header_footer = False
            sec.footer.is_linked_to_previous = False
            _docx_pagefield(sec.footer)
            _docx_pgnumtype(sec, "decimal", start=1)
            ctx["body_section_started"] = True
            page_break = False  # the section break already starts a new page
        else:
            # Each chapter starts on a new page (mirrors LaTeX \chapter).
            page_break = tw and level == 1
        if numbered_chapter:
            # New chapter: 「第N章　標題」, centered; reset figure/table/section counters.
            ctx["chapter"] = ctx.get("chapter", 0) + 1
            ctx["fig"] = 0
            ctx["tab"] = 0
            ctx["sec"] = 0
            ctx["subsec"] = 0
            p = docx.add_heading("", level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"第{_cjk_num(ctx['chapter'])}章　")
            _add_inline_docx(p, _children(block), style, order)
        elif tw and level == 1:
            # Unnumbered chapter (參考文獻/附錄/誌謝…): centered, no 第N章.
            p = docx.add_heading("", level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_docx(p, _children(block), style, order)
        elif tw and level in (2, 3):
            # Number sections to match the LaTeX/PDF: h2→\section (1.1, ctex
            # centers it), h3→\subsection (1.1.1, flush left). Without this the
            # DOCX showed bare unnumbered headings while the PDF was numbered.
            chap = ctx.get("chapter", 0)
            if level == 2:
                ctx["sec"] = ctx.get("sec", 0) + 1
                ctx["subsec"] = 0
                num = f"{chap}.{ctx['sec']}"
            else:
                ctx["subsec"] = ctx.get("subsec", 0) + 1
                num = f"{chap}.{ctx.get('sec', 0)}.{ctx['subsec']}"
            p = docx.add_heading("", level=level)
            if level == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{num}　")
            _add_inline_docx(p, _children(block), style, order)
        else:
            p = docx.add_heading("", level=min(max(level, 1), 4))
            _add_inline_docx(p, _children(block), style, order)
        if tw:
            _style_heading_runs(p)
            if page_break:
                p.paragraph_format.page_break_before = True
    elif t in ("bulletList", "orderedList"):
        list_style = "List Bullet" if t == "bulletList" else "List Number"
        for item in _children(block):
            for child in _children(item):  # each listItem wraps paragraph(s)
                p = docx.add_paragraph(style=list_style)
                _add_inline_docx(p, _children(child), style, order)
    elif t == "blockquote":
        for child in _children(block):
            p = docx.add_paragraph(style="Quote")
            _add_inline_docx(p, _children(child), style, order)
    elif t == "codeBlock":
        p = docx.add_paragraph()
        run = p.add_run(_block_text(block))
        run.font.name = "Courier New"
    elif t == "horizontalRule":
        docx.add_paragraph("―" * 20)
    elif t == "figure":
        attrs = block.get("attrs") or {}
        cap = attrs.get("caption") or ""
        img = _image_bytes(attrs.get("src", ""))
        if img:
            try:
                docx.add_picture(io.BytesIO(img[0]), width=Inches(5.5))
                docx.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:  # noqa: BLE001 — unsupported image → caption only
                img = None
        # Caption below the figure; chapter-scoped 「圖 2-1」 numbering (only
        # captioned figures are numbered, mirroring the editor convention).
        cap_text = cap or ("" if img else "[image]")
        if cap and tw:
            cap_text = f"圖 {_docx_caption_num(ctx, 'fig')}　{cap}"
        capp = docx.add_paragraph()
        capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = capp.add_run(cap_text)
        run.italic = not tw
        if tw:
            _apply_run_font(run)
    elif t == "figureList":
        pass  # an editor-only live aid; the figures themselves render inline
    elif t == "mathBlock":
        p = docx.add_paragraph()
        p.add_run(f"$${(block.get('attrs') or {}).get('latex', '')}$$")
    elif t == "tableBlock":
        caption, rows = _table_parts(block)
        if caption:
            # Caption above the table; chapter-scoped 「表 2-1」 numbering.
            cap_text = (
                f"表 {_docx_caption_num(ctx, 'tab')}　{caption}" if tw else caption
            )
            cp = docx.add_paragraph()
            if tw:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(cap_text)
            run.italic = not tw
            if tw:
                _apply_run_font(run)
        if rows:
            ncols = max(len(r) for r in rows)
            tbl = docx.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Table Grid"
            for r, cells in enumerate(rows):
                for c in range(ncols):
                    tbl.cell(r, c).text = cells[c] if c < len(cells) else ""
    elif t == "tableList":
        pass  # editor-only live aid; the tables render inline
    elif t == "referenceList":
        # Render the reference list right here (mirrors the LaTeX path), so the
        # bibliography lands at the node's position in all three formats.
        ctx["refs_inline"] = True
        resolved = ctx.get("resolved") or []
        if resolved:
            if ctx.get("tw"):
                _docx_refs_grouped(docx, resolved, ctx["refs_label"], ctx["locale"])
            else:
                docx.add_heading(ctx["refs_label"], level=1)
                # Number by first-appearance order (same universe as in-text
                # labels and the LaTeX list) — enumerate(resolved) drifted as
                # soon as an unlinked citation occupied an in-text number.
                for c in resolved:
                    docx.add_paragraph(
                        full_reference(
                            c,
                            ctx["style"],
                            _citation_number(ctx["order"], _cite_key(c)),
                        )
                    )
    else:  # paragraph and any unknown block → a plain paragraph of its inline content
        p = docx.add_paragraph()
        if tw:
            # 首行縮排兩個字 — the Chinese-thesis convention ctex applies by default
            # (\parindent = 2 字寬 ≈ 2×body-size); without it Word body paragraphs
            # sat flush-left while the PDF indented them.
            p.paragraph_format.first_line_indent = Pt(28)
        _add_inline_docx(p, _children(block), style, order)


def _set_academic_font(docx: Document, size_pt: int = 12) -> None:
    """Use a paper-like serif (Times New Roman + 標楷體 for CJK) for the body,
    at the given point size (14pt for twthesis, the Taiwan-thesis norm)."""
    try:
        normal = docx.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(size_pt)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "標楷體")
    except Exception:  # noqa: BLE001 — styling is best-effort
        pass


def _docx_page_layout(docx: Document) -> None:
    """Taiwan-thesis page setup: 3cm margins + 1.5 line spacing."""
    for sec in docx.sections:
        sec.top_margin = Cm(3)
        sec.bottom_margin = Cm(3)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(3)
    pf = docx.styles["Normal"].paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def _docx_pagefield(footer) -> None:
    """Put a centered PAGE field in a section's footer."""
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    txt = OxmlElement("w:t")
    txt.text = "1"
    r.append(txt)
    fld.append(r)
    p._p.append(fld)


def _docx_pgnumtype(section, fmt: str, start: int | None = None) -> None:
    """Set a section's page-number format (lowerRoman / decimal) and optional
    restart. pgNumType must sit just before w:cols in the schema order."""
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    cols = sectPr.find(qn("w:cols"))
    if cols is not None:
        cols.addprevious(pg)
    else:
        sectPr.append(pg)


def _docx_section_margins(section) -> None:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)


def _has_numbered_chapter(content: dict) -> bool:
    for b in _children(content):
        if b.get("type") == "heading" and (b.get("attrs") or {}).get("level") == 1:
            if not _is_unnumbered_chapter(_plain_text(b)):
                return True
    return False


def _docx_cover(docx: Document, cover: dict, fallback_title: str) -> None:
    """Bilingual Taiwan-thesis title page (mirrors _tw_cover_latex)."""

    def g(k: str) -> str:
        return (cover.get(k) or "").strip()

    def line(text: str, size: int, bold: bool = False, after: int = 6) -> None:
        p = docx.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        _apply_run_font(r)

    line("", 12)
    head_zh = " ".join(x for x in (g("universityZh"), g("departmentZh")) if x)
    if head_zh:
        line(head_zh, 18)
    if g("degreeZh"):
        line(g("degreeZh"), 18, after=12)
    for x in (g("departmentEn"), g("universityEn"), g("degreeEn")):
        if x:
            line(x, 13)
    line("", 18)
    line(g("titleZh") or fallback_title, 16, after=10)
    if g("titleEn"):
        line(g("titleEn"), 13)
    line("", 24)
    if g("advisorZh"):
        line(f"指導教授：{g('advisorZh')}", 14)
    if g("advisorEn"):
        line(f"Advisor: {g('advisorEn')}", 13, after=12)
    if g("studentZh"):
        line(f"研究生：{g('studentZh')}", 14)
    if g("studentEn"):
        line(f"Graduate Student: {g('studentEn')}", 13)
    line("", 24)
    if g("dateZh"):
        line(g("dateZh"), 14)
    if g("dateEn"):
        line(g("dateEn"), 13)


def _docx_refs_grouped(
    docx: Document, resolved: list[dict], refs_label: str, locale: str
) -> None:
    """Grouped (網路資源/中文/英文) reference list with bold sub-headers and a
    hanging indent — the Taiwan-thesis convention (mirrors _render_refs_latex)."""
    h = docx.add_heading(refs_label, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.page_break_before = True  # 參考文獻 starts a new page
    _style_heading_runs(h)
    labels = _REF_GROUP_LABELS.get(locale, _REF_GROUP_LABELS["en"])
    groups: dict[str, list[dict]] = {"web": [], "zh": [], "en": []}
    for c in resolved:
        groups[_ref_group(c)].append(c)
    for key in ("web", "zh", "en"):
        g = sorted(groups[key], key=_ref_sort_key)
        if not g:
            continue
        hp = docx.add_paragraph()
        hr = hp.add_run(labels[key])
        hr.bold = True
        _apply_run_font(hr)
        for c in g:
            p = docx.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Pt(28)  # 2 字寬, matches the LaTeX \hangindent=2em
            pf.first_line_indent = Pt(-28)  # hanging indent
            _apply_run_font(p.add_run(_tw_reference(c)))


def to_docx(
    document: dict,
    style: str,
    refs_label: str,
    template: str = "article",
    locale: str = "zh-Hant",
    cover: dict | None = None,
) -> bytes:
    content = _prepared_content(document)
    title = document.get("title") or "(untitled)"
    citations = collect_citations(content)
    order = [_cite_key(c) for c in citations]
    tw = template == "twthesis"

    docx = Document()
    _set_academic_font(docx, 14 if tw else 12)
    cover_active = bool(tw and cover and any((v or "").strip() for v in cover.values()))
    # Front matter (cover) gets Roman page numbers; the body flips to Arabic at
    # the first numbered chapter via a section break — but only when there's a
    # cover to anchor the unnumbered first page.
    roman_split = bool(cover_active and _has_numbered_chapter(content))
    if tw:
        _docx_page_layout(docx)
        _docx_pagefield(docx.sections[0].footer)
        if cover_active:
            docx.sections[0].different_first_page_header_footer = True
            if roman_split:
                _docx_pgnumtype(docx.sections[0], "lowerRoman")
            _docx_cover(docx, cover, title)
            # The page break to the next page is deferred: the first numbered
            # chapter starts a new *section* (Roman→Arabic), which also breaks the
            # page — so we don't add a manual break here (it would blank a page).
        else:
            docx.add_heading(title, level=0)
    else:
        docx.add_heading(title, level=0)

    resolved = [c for c in citations if not c.get("unlinked")]
    ctx = {
        "tw": tw,
        "chapter": 0,
        "fig": 0,
        "tab": 0,
        "resolved": resolved,
        "order": order,
        "refs_label": refs_label,
        "locale": locale,
        "style": style,
        "refs_inline": False,
        "roman_split": roman_split,
        "body_section_started": False,
        # The cover needs the following content on a fresh page; consumed by the
        # first block (a section break if it's a chapter, else a page break).
        "pending_cover_break": cover_active,
    }
    for block in _children(content):
        _render_block_docx(docx, block, style, order, ctx)

    # Append the reference list only if no referenceList node placed it inline.
    if resolved and not ctx["refs_inline"]:
        if tw:
            _docx_refs_grouped(docx, resolved, refs_label, locale)
        else:
            docx.add_heading(refs_label, level=1)
            for c in resolved:
                docx.add_paragraph(
                    full_reference(c, style, _citation_number(order, _cite_key(c)))
                )

    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()


# ---------- LaTeX ----------

_LATEX_ESCAPE = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _esc(s: str) -> str:
    return "".join(_LATEX_ESCAPE.get(ch, ch) for ch in (s or ""))


def _inline_latex(nodes: list[dict], style: str, order: list[str]) -> str:
    parts: list[str] = []
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            s = _esc(n.get("text", ""))
            marks = {m.get("type") for m in (n.get("marks") or [])}
            if "code" in marks:
                s = r"\texttt{" + s + "}"
            if "strike" in marks:
                s = r"\sout{" + s + "}"
            if "italic" in marks:
                s = r"\textit{" + s + "}"
            if "bold" in marks:
                s = r"\textbf{" + s + "}"
            parts.append(s)
        elif t == "citation":
            a = n.get("attrs") or {}
            num = _citation_number(order, _cite_key(a))
            parts.append(_esc(in_text_label(a, style, num)))
        elif t == "mathInline":
            # LaTeX source goes through verbatim — do NOT escape.
            parts.append(f"${(n.get('attrs') or {}).get('latex', '')}$")
    return "".join(parts)


def _render_block_latex(block: dict, style: str, order: list[str], ctx: dict) -> str:
    t = block.get("type")
    if t == "heading":
        level = (block.get("attrs") or {}).get("level", 1)
        prefix = ""
        if ctx.get("chapters"):
            # Thesis layout (report class): top-level headings are chapters —
            # ctex renders them as 「第一章 緒論」.
            cmd = {1: "chapter", 2: "section", 3: "subsection"}.get(level, "paragraph")
            if level == 1:
                if _is_unnumbered_chapter(_plain_text(block)):
                    cmd = "chapter*"  # 參考文獻/附錄/誌謝… — no 第N章 number
                elif not ctx.get("arabic_started"):
                    # Front matter (摘要/目錄/誌謝) is Roman; the body switches to
                    # Arabic at the first *numbered* chapter — the Taiwan norm.
                    prefix = "\\pagenumbering{arabic}\n"
                    ctx["arabic_started"] = True
        else:
            cmd = {1: "section", 2: "subsection", 3: "subsubsection"}.get(
                level, "paragraph"
            )
        return (
            prefix + f"\\{cmd}{{{_inline_latex(_children(block), style, order)}}}\n\n"
        )
    if t == "referenceList":
        # The doc carries an explicit 參考文獻清單 node — render the reference
        # list right here (grouped, Taiwan format) instead of auto-appending it.
        ctx["refs_inline"] = True
        return ctx.get("refs_latex", "")
    if t in ("bulletList", "orderedList"):
        env = "itemize" if t == "bulletList" else "enumerate"
        lines = [f"\\begin{{{env}}}"]
        for item in _children(block):
            inner = "".join(
                _inline_latex(_children(child), style, order)
                for child in _children(item)
            )
            lines.append(f"  \\item {inner}")
        lines.append(f"\\end{{{env}}}")
        return "\n".join(lines) + "\n\n"
    if t == "blockquote":
        inner = "".join(
            _inline_latex(_children(child), style, order) + "\n"
            for child in _children(block)
        )
        return f"\\begin{{quote}}\n{inner}\\end{{quote}}\n\n"
    if t == "codeBlock":
        return f"\\begin{{verbatim}}\n{_block_text(block)}\n\\end{{verbatim}}\n\n"
    if t == "horizontalRule":
        return "\\hrulefill\n\n"
    if t == "figure":
        attrs = block.get("attrs") or {}
        cap = attrs.get("caption") or ""
        img = _image_bytes(attrs.get("src", ""))
        if img:
            data, ext = img
            ctx["fig_n"] += 1
            fname = f"fig{ctx['fig_n']}.{ext}"
            ctx["images"].append((fname, data))
            capline = f"\\caption{{{_esc(cap)}}}\n" if cap else ""
            # float's [H] = put it HERE. Plain [h] is only a hint — when LaTeX
            # can't honor it the figure floats away and piles up at the end of
            # the document, which reads as "my images are all misplaced".
            return (
                "\\begin{figure}[H]\n\\centering\n"
                f"\\includegraphics[width=0.8\\linewidth]{{{fname}}}\n{capline}\\end{{figure}}\n\n"
            )
        return f"\\textit{{[{_esc(cap or 'image')}]}}\n\n"
    if t == "figureList":
        return ""  # an editor-only live aid; the figures render inline
    if t == "mathBlock":
        return f"\\[{(block.get('attrs') or {}).get('latex', '')}\\]\n\n"
    if t == "tableBlock":
        caption, rows = _table_parts(block)
        if not rows:
            return ""
        ncols = max(len(r) for r in rows)
        colspec = "|" + "X|" * ncols
        # Page-breakable tables break only BETWEEN rows, so a transcript-length
        # cell (pages of text in one cell) must first be split into
        # continuation rows — otherwise the row overflows the page no matter
        # which environment renders it.
        chunk_limit = _cell_chunk_chars(ncols)
        phys_rows: list[tuple[list[str], bool]] = []  # (cells, is_continuation)
        for r in rows:
            chunk_lists = [
                _chunk_cell(r[c] if c < len(r) else "", chunk_limit)
                for c in range(ncols)
            ]
            for idx in range(max(len(cl) for cl in chunk_lists)):
                cells = [cl[idx] if idx < len(cl) else "" for cl in chunk_lists]
                phys_rows.append((cells, idx > 0))
        lines = []
        for cells, is_cont in phys_rows:
            if not is_cont and lines:
                lines.append("\\hline")
            lines.append(" & ".join(_esc(c) for c in cells) + " \\\\")
        body = "\n".join(lines)
        if ctx.get("longtable", True):
            # xltabular = longtable + tabularx: X columns wrap within
            # \linewidth AND the table breaks across pages between rows.
            cap = f"\\caption{{{_esc(caption)}}}\\\\\n" if caption else ""
            return (
                f"\\begin{{xltabular}}{{\\linewidth}}{{{colspec}}}\n"
                f"{cap}\\hline\n{body}\n\\hline\n\\end{{xltabular}}\n\n"
            )
        # longtable can't live inside twocolumn layouts (IEEE / twocolumn
        # article) — fall back to a fixed tabularx anchored in place.
        cap = f"\\caption{{{_esc(caption)}}}\n" if caption else ""
        return (
            "\\begin{table}[H]\n\\centering\n"
            f"\\begin{{tabularx}}{{\\linewidth}}{{{colspec}}}\n\\hline\n"
            f"{body}\n\\hline\n\\end{{tabularx}}\n{cap}\\end{{table}}\n\n"
        )
    if t == "tableList":
        return ""  # editor-only live aid; the tables render inline
    return _inline_latex(_children(block), style, order) + "\n\n"


# Document layouts → the \documentclass line. "twthesis" approximates the
# common Taiwan graduate-thesis conventions: single-column A4 report with
# chapter-level headings, 3cm margins and 1.5 line spacing.
_LATEX_TEMPLATES = {
    "article": "\\documentclass{article}",
    "twocolumn": "\\documentclass[twocolumn]{article}",
    "ieee": "\\documentclass[conference]{IEEEtran}",
    "twthesis": "\\documentclass[a4paper,12pt]{report}",
}

_SENTENCE_RE = re.compile(r"[^。！？!?]*(?:[。！？!?]|$)")


def _cell_chunk_chars(ncols: int) -> int:
    """Chunk size for one physical table row, scaled by column width.

    A full \\linewidth fits ~40 CJK chars per line; an X column gets 1/ncols of
    that. Aim for ~6 lines per chunk: a chunk that doesn't fit at the page
    bottom leaves its whole height blank, so finer chunks = tighter pages."""
    return max(60, 240 // max(ncols, 1))


def _chunk_cell(text: str, limit: int) -> list[str]:
    """Split oversized cell text into sentence-aligned chunks (each rendered as
    a continuation row). Short cells come back as a single chunk."""
    if len(text) <= limit:
        return [text]
    chunks: list[str] = []
    cur = ""
    for sent in _SENTENCE_RE.findall(text):
        if not sent:
            continue
        if cur and len(cur) + len(sent) > limit:
            chunks.append(cur)
            cur = ""
        # a single sentence longer than the limit gets hard-split
        while len(sent) > limit:
            chunks.append(sent[:limit])
            sent = sent[limit:]
        cur += sent
    if cur:
        chunks.append(cur)
    return chunks or [""]


def _tw_cover_latex(cover: dict, fallback_title: str) -> str:
    """A Taiwan-thesis bilingual title page (titlepage env). Mirrors the common
    layout: university/department/degree (zh then en) → bilingual title →
    advisor → student → ROC date. Empty fields are skipped; the Chinese title
    falls back to the document title."""

    def g(k: str) -> str:
        return (cover.get(k) or "").strip()

    out = ["\\begin{titlepage}\n\\centering\n\\thispagestyle{empty}\n\\vspace*{2em}\n"]
    head_zh = " ".join(x for x in (g("universityZh"), g("departmentZh")) if x)
    if head_zh:
        out.append(f"{{\\zihao{{2}} {_esc(head_zh)}\\par}}\n\\vspace{{0.4em}}\n")
    if g("degreeZh"):
        out.append(f"{{\\zihao{{2}} {_esc(g('degreeZh'))}\\par}}\n\\vspace{{0.8em}}\n")
    head_en = [x for x in (g("departmentEn"), g("universityEn"), g("degreeEn")) if x]
    if head_en:
        out.append("{\\large " + "\\\\\n".join(_esc(x) for x in head_en) + "\\par}\n")
    out.append("\\vspace{2.5em}\n")
    title_zh = g("titleZh") or fallback_title
    if title_zh:
        out.append(f"{{\\zihao{{3}} {_esc(title_zh)}\\par}}\n\\vspace{{0.8em}}\n")
    if g("titleEn"):
        out.append(f"{{\\large {_esc(g('titleEn'))}\\par}}\n")
    out.append("\\vspace{3em}\n")
    if g("advisorZh"):
        out.append(f"{{\\zihao{{4}} 指導教授：{_esc(g('advisorZh'))}\\par}}\n")
    if g("advisorEn"):
        out.append(f"{{\\large Advisor: {_esc(g('advisorEn'))}\\par}}\n")
    out.append("\\vspace{1.5em}\n")
    if g("studentZh"):
        out.append(f"{{\\zihao{{4}} 研究生：{_esc(g('studentZh'))}\\par}}\n")
    if g("studentEn"):
        out.append(f"{{\\large Graduate Student: {_esc(g('studentEn'))}\\par}}\n")
    out.append("\\vfill\n")
    if g("dateZh"):
        out.append(f"{{\\zihao{{4}} {_esc(g('dateZh'))}\\par}}\n")
    if g("dateEn"):
        out.append(f"{{\\large {_esc(g('dateEn'))}\\par}}\n")
    out.append("\\end{titlepage}\n")
    return "".join(out)


# Bold sub-headers for the Taiwan-thesis three-way reference split.
_REF_GROUP_LABELS = {
    "zh-Hant": {"web": "網路資源", "zh": "中文", "en": "英文"},
    "en": {"web": "Web Resources", "zh": "Chinese", "en": "English"},
}


def _render_refs_latex(
    citations: list[dict],
    style: str,
    order: list[str],
    refs_label: str,
    ctx: dict,
    locale: str,
    template: str,
) -> str:
    """Reference list as LaTeX. For twthesis: grouped (網路資源 / 中文 / 英文)
    with bold sub-headers, full-width Chinese punctuation and hanging indent —
    the Taiwan-thesis convention. Other templates keep the flat numbered
    itemize. Unlinked (待補來源) citations are dropped — only resolved sources
    appear in the list, mirroring the editor's referenceList node."""
    resolved = [c for c in citations if not c.get("unlinked")]
    if not resolved:
        return ""
    refs_cmd = "chapter*" if ctx.get("chapters") else "section*"
    head = f"\\{refs_cmd}{{{_esc(refs_label)}}}\n"
    if template != "twthesis":
        items = "\n".join(
            f"  \\item {_esc(full_reference(c, style, _citation_number(order, _cite_key(c))))}"
            for c in resolved
        )
        return head + f"\\begin{{itemize}}\n{items}\n\\end{{itemize}}\n\n"
    labels = _REF_GROUP_LABELS.get(locale, _REF_GROUP_LABELS["en"])
    groups: dict[str, list[dict]] = {"web": [], "zh": [], "en": []}
    for c in resolved:
        groups[_ref_group(c)].append(c)
    out = [head, "{\\setlength{\\parindent}{0pt}\n"]
    for key in ("web", "zh", "en"):
        g = sorted(groups[key], key=_ref_sort_key)
        if not g:
            continue
        out.append(
            f"\\par\\medskip\\noindent\\textbf{{{_esc(labels[key])}}}\\par\\smallskip\n"
        )
        for c in g:
            out.append(
                f"\\noindent\\hangindent=2em\\hangafter=1\\relax {_tw_reference_latex(c)}"
                "\\par\\vspace{4pt}\n"
            )
    out.append("}\n\n")
    return "".join(out)


def to_latex(
    document: dict,
    style: str,
    refs_label: str,
    template: str = "article",
    locale: str = "zh-Hant",
    cover: dict | None = None,
) -> tuple[str, list[tuple[str, bytes]]]:
    """Render to LaTeX. Returns (tex_source, images) where images is
    [(filename, bytes)] referenced by \\includegraphics — the route bundles them
    into a .zip when non-empty. `template` picks the \\documentclass."""
    content = _prepared_content(document)
    title = document.get("title") or "(untitled)"
    citations = collect_citations(content)
    order = [_cite_key(c) for c in citations]

    # longtable (inside xltabular) can't be used in twocolumn layouts.
    ctx: dict = {
        "images": [],
        "fig_n": 0,
        "longtable": template not in ("twocolumn", "ieee"),
        "chapters": template == "twthesis",
        "refs_inline": False,
    }
    # Precompute the reference list so an inline `referenceList` node can emit it
    # at its own position; otherwise it's appended after the body.
    refs_latex = _render_refs_latex(
        citations, style, order, refs_label, ctx, locale, template
    )
    ctx["refs_latex"] = refs_latex
    body = "".join(
        _render_block_latex(b, style, order, ctx) for b in _children(content)
    )
    refs = "" if ctx.get("refs_inline") else refs_latex

    docclass = _LATEX_TEMPLATES.get(template, _LATEX_TEMPLATES["article"])
    # ctex → CJK; graphicx → figures; ulem → \sout; hyperref → links.
    # ctex/fontspec need XeLaTeX/LuaLaTeX — pdfLaTeX fails with "fontset fandol
    # unavailable" — so only emit them when the content actually contains CJK;
    # a pure-Latin document then compiles untouched on Overleaf's default
    # pdfLaTeX. The magic comment makes TeXShop / VS Code / local latexmk pick
    # XeLaTeX automatically (Overleaf ignores it — switch its compiler menu).
    # A Taiwan-thesis cover replaces \maketitle with a bilingual title page.
    cover_active = bool(
        template == "twthesis"
        and cover
        and any((v or "").strip() for v in cover.values())
    )
    cover_text = " ".join(cover.values()) if cover_active else ""
    needs_cjk = bool(_CJK_RE.search(title + body + refs + cover_text))
    magic = "% !TEX program = xelatex\n" if needs_cjk else ""
    # ctex's default Fandol font is missing many Traditional-Chinese-only
    # glyphs (為/致/夠/謹…), which render as boxes. Prefer Noto Serif CJK TC
    # (full coverage, bundled on Overleaf); keep Fandol where Noto is absent
    # so local compiles don't break with a "font not found" error.
    # CJK body font. Taiwan theses are set in 標楷體 (a kai/楷 brush style) — use
    # the bundled "AR PL KaitiM Big5" (kai) for twthesis; other templates keep the
    # neutral Noto serif. Both fall back to nothing if the font is absent (local /
    # Overleaf compiles then use ctex's default rather than erroring).
    if template == "twthesis":
        cjk_main = (
            "\\IfFontExistsTF{AR PL KaitiM Big5}{\\setCJKmainfont{AR PL KaitiM Big5}}"
            "{\\IfFontExistsTF{Noto Serif CJK TC}{\\setCJKmainfont{Noto Serif CJK TC}}{}}\n"
        )
    else:
        cjk_main = (
            "\\IfFontExistsTF{Noto Serif CJK TC}"
            "{\\setCJKmainfont{Noto Serif CJK TC}}{}\n"
        )
    if not needs_cjk:
        cjk_packages = ""
    elif locale == "zh-Hant":
        # heading=true localizes the heading machinery (thesis layout:
        # \chapter → 「第一章」 instead of "Chapter 1"). ctex defaults to
        # Simplified-Chinese caption labels (图/表) — pin the Traditional forms.
        ctex_opts = "[heading=true]" if template == "twthesis" else ""
        cjk_packages = (
            f"\\usepackage{ctex_opts}{{ctex}}\n"
            "\\ctexset{figurename={圖},tablename={表}}\n" + cjk_main
        )
    else:
        # English doc that merely contains CJK: typeset the CJK but keep
        # English caption labels (Figure/Table) and Western layout.
        cjk_packages = "\\usepackage[scheme=plain]{ctex}\n" + cjk_main
    # Taiwan-thesis conventions: 3cm margins on A4 + 1.5 line spacing.
    layout = (
        "\\usepackage[a4paper,margin=3cm]{geometry}\n"
        "\\usepackage{setspace}\n\\onehalfspacing\n"
        if template == "twthesis"
        else ""
    )
    # twthesis Western body font (Times New Roman → Liberation Serif, a metric-
    # compatible clone) + chapter-scoped figure/table numbering (圖 1-1, 表 2-3).
    # Needs fontspec, which ctex already loaded above, so it follows cjk_packages.
    tw_extras = (
        (
            "\\IfFontExistsTF{Times New Roman}{\\setmainfont{Times New Roman}}"
            "{\\IfFontExistsTF{Liberation Serif}{\\setmainfont{Liberation Serif}}{}}\n"
            "\\renewcommand{\\thefigure}{\\thechapter-\\arabic{figure}}\n"
            "\\renewcommand{\\thetable}{\\thechapter-\\arabic{table}}\n"
            # Caption label separator = a space (「圖 2-1 標題」), not the default
            # colon — matches the DOCX output and the Taiwan-thesis convention.
            "\\usepackage{caption}\n\\captionsetup{labelsep=quad}\n"
        )
        if template == "twthesis" and needs_cjk
        else ""
    )
    # 14pt body for twthesis. report tops out at 12pt, so set the size from ctex's
    # \zihao (四號 = 14pt) once the document opens; headings keep the class sizes.
    # \url in the reference list should read in the body font (not typewriter)
    # while still breaking long DOIs/URLs across lines.
    url_style = "\\urlstyle{same}\n" if template == "twthesis" else ""
    body_size = "\\zihao{4}\n" if template == "twthesis" else ""
    # Taiwan thesis: front matter (摘要/目錄/誌謝) in Roman numerals; the body
    # flips to Arabic at the first chapter (injected in _render_block_latex).
    front_matter = "\\pagenumbering{roman}\n" if template == "twthesis" else ""
    title_block = _tw_cover_latex(cover, title) if cover_active else "\\maketitle\n"
    tex = (
        magic + f"{docclass}\n" + cjk_packages + layout + "\\usepackage{graphicx}\n"
        "\\usepackage{float}\n"  # [H] placement for figures/tables
        "\\usepackage{xltabular}\n"  # width-bounded wrapping columns + page-breakable tables
        "\\usepackage[normalem]{ulem}\n"
        "\\usepackage{hyperref}\n"
        + url_style
        + tw_extras
        + f"\\title{{{_esc(title)}}}\n"
        "\\author{}\n\\date{}\n"
        "\\begin{document}\n"
        + body_size
        + title_block
        + front_matter
        + "\n"
        + body
        + refs
        + "\\end{document}\n"
    )
    return tex, ctx["images"]


# ---------- Markdown / plain text / HTML (preview + browser-print PDF) ----------


def _inline_md(nodes: list[dict], style: str, order: list[str]) -> str:
    parts: list[str] = []
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            s = n.get("text", "")
            marks = {m.get("type") for m in (n.get("marks") or [])}
            if "code" in marks:
                s = f"`{s}`"
            if "strike" in marks:
                s = f"~~{s}~~"
            if "italic" in marks:
                s = f"*{s}*"
            if "bold" in marks:
                s = f"**{s}**"
            parts.append(s)
        elif t == "citation":
            a = n.get("attrs") or {}
            parts.append(in_text_label(a, style, _citation_number(order, _cite_key(a))))
        elif t == "mathInline":
            parts.append(f"${(n.get('attrs') or {}).get('latex', '')}$")
    return "".join(parts)


def to_markdown(document: dict, style: str, refs_label: str) -> str:
    content = _prepared_content(document)
    title = document.get("title") or "(untitled)"
    citations = collect_citations(content)
    order = [_cite_key(c) for c in citations]
    out = [f"# {title}", ""]
    for b in _children(content):
        t = b.get("type")
        if t == "heading":
            lvl = (b.get("attrs") or {}).get("level", 1)
            out.append(
                "#" * min(lvl + 1, 6) + " " + _inline_md(_children(b), style, order)
            )
        elif t in ("bulletList", "orderedList"):
            for i, item in enumerate(_children(b), 1):
                inner = "".join(
                    _inline_md(_children(ch), style, order) for ch in _children(item)
                )
                out.append(("- " if t == "bulletList" else f"{i}. ") + inner)
        elif t == "blockquote":
            for ch in _children(b):
                out.append("> " + _inline_md(_children(ch), style, order))
        elif t == "codeBlock":
            out.append("```\n" + _block_text(b) + "\n```")
        elif t == "horizontalRule":
            out.append("---")
        elif t == "figure":
            a = b.get("attrs") or {}
            out.append(f"![{a.get('caption', '')}]({a.get('src', '')})")
            if a.get("caption"):
                out.append(f"*{a['caption']}*")
        elif t == "mathBlock":
            out.append(f"$$\n{(b.get('attrs') or {}).get('latex', '')}\n$$")
        elif t == "tableBlock":
            cap, rows = _table_parts(b)
            if cap:
                out.append(f"**{cap}**")
            if rows:
                ncols = max(len(r) for r in rows)
                out.append("| " + " | ".join((rows[0] + [""] * ncols)[:ncols]) + " |")
                out.append("| " + " | ".join(["---"] * ncols) + " |")
                for r in rows[1:]:
                    out.append("| " + " | ".join((r + [""] * ncols)[:ncols]) + " |")
        elif t in ("tableList", "referenceList"):
            continue
        else:
            out.append(_inline_md(_children(b), style, order))
        out.append("")
    resolved = [c for c in citations if not c.get("unlinked")]
    if resolved:
        out.append(f"## {refs_label}")
        out.append("")
        # List marker stays sequential (markdown renderers renumber anyway);
        # the reference's own [n] must match the in-text numbering universe.
        for i, c in enumerate(resolved):
            out.append(
                f"{i + 1}. {full_reference(c, style, _citation_number(order, _cite_key(c)))}"
            )
    return "\n".join(out).strip() + "\n"


def to_text(document: dict, style: str, refs_label: str) -> str:
    """Strip-to-plain-text: citations as in-text markers, references appended."""
    content = _prepared_content(document)
    title = document.get("title") or "(untitled)"
    citations = collect_citations(content)
    order = [_cite_key(c) for c in citations]

    def inline(nodes: list[dict]) -> str:
        parts: list[str] = []
        for n in nodes or []:
            t = n.get("type")
            if t == "text":
                parts.append(n.get("text", ""))
            elif t == "citation":
                a = n.get("attrs") or {}
                parts.append(
                    in_text_label(a, style, _citation_number(order, _cite_key(a)))
                )
            elif t == "mathInline":
                parts.append(f"${(n.get('attrs') or {}).get('latex', '')}$")
        return "".join(parts)

    out = [title, ""]
    for b in _children(content):
        t = b.get("type")
        if t in ("bulletList", "orderedList"):
            for i, item in enumerate(_children(b), 1):
                inner = "".join(inline(_children(ch)) for ch in _children(item))
                out.append(("- " if t == "bulletList" else f"{i}. ") + inner)
        elif t == "codeBlock":
            out.append(_block_text(b))
        elif t == "figure":
            cap = (b.get("attrs") or {}).get("caption") or ""
            out.append(f"[{('圖: ' + cap) if cap else 'image'}]")
        elif t == "mathBlock":
            out.append(f"${(b.get('attrs') or {}).get('latex', '')}$")
        elif t == "tableBlock":
            cap, rows = _table_parts(b)
            if cap:
                out.append(cap)
            for r in rows:
                out.append("\t".join(r))
        elif t in ("tableList", "referenceList"):
            continue
        else:
            out.append(inline(_children(b)))
        out.append("")
    resolved = [c for c in citations if not c.get("unlinked")]
    if resolved:
        out.append(refs_label)
        for c in resolved:
            num = _citation_number(order, _cite_key(c))
            out.append(f"[{num}] {full_reference(c, style, num)}")
    return "\n".join(out).strip() + "\n"


_HTML_ESCAPE = {"&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;"}


def _h(s: str) -> str:
    return "".join(_HTML_ESCAPE.get(ch, ch) for ch in (s or ""))


def _inline_html(nodes: list[dict], style: str, order: list[str]) -> str:
    parts: list[str] = []
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            s = _h(n.get("text", ""))
            marks = {m.get("type") for m in (n.get("marks") or [])}
            if "code" in marks:
                s = f"<code>{s}</code>"
            if "strike" in marks:
                s = f"<s>{s}</s>"
            if "italic" in marks:
                s = f"<em>{s}</em>"
            if "bold" in marks:
                s = f"<strong>{s}</strong>"
            parts.append(s)
        elif t == "citation":
            a = n.get("attrs") or {}
            parts.append(
                _h(in_text_label(a, style, _citation_number(order, _cite_key(a))))
            )
        elif t == "mathInline":
            # _h(): raw latex straight into HTML is stored XSS (docs are shared
            # globally). MathJax parses the entity-escaped text just fine.
            parts.append("\\(" + _h((n.get("attrs") or {}).get("latex", "")) + "\\)")
    return "".join(parts)


# Academic, paper-like styling: serif body, justified text, numbered refs.
_HTML_CSS = """
:root { color-scheme: light; }
body { font-family: "Noto Serif TC","Noto Serif CJK TC","Songti TC","Times New Roman",Georgia,serif;
  max-width: 760px; margin: 2.5rem auto; padding: 0 1.5rem; line-height: 1.7;
  color: #1a1a1a; background: #fff; text-align: justify; }
h1 { text-align: center; font-size: 1.7rem; margin-bottom: 1.6rem; }
h2 { font-size: 1.3rem; margin-top: 1.8rem; border-bottom: 1px solid #ddd; padding-bottom: .2rem; }
h3 { font-size: 1.1rem; margin-top: 1.4rem; }
figure { text-align: center; margin: 1.4rem 0; }
figure img { max-width: 100%; }
figcaption { font-size: .9rem; color: #555; margin-top: .4rem; }
table { border-collapse: collapse; width: 100%; margin: 1rem 0; }
th, td { border: 1px solid #999; padding: .4rem .6rem; }
blockquote { border-left: 3px solid #ccc; margin: 1rem 0; padding-left: 1rem; color: #444; }
pre { background: #f5f5f5; padding: .8rem; overflow-x: auto; font-family: ui-monospace,monospace; }
ol.refs li { margin-bottom: .4rem; }
@media print { body { margin: 0; max-width: none; } }
"""

# Taiwan-thesis preview: an A4 "page" on a gray desk, 標楷體 (kai) for CJK +
# Times for Latin, centered 第N章 headings, hanging-indent grouped references —
# a screen mirror of the XeLaTeX PDF so the preview matches the chosen layout.
_HTML_CSS_TW = """
:root { color-scheme: light; }
body { margin: 0; padding: 2rem 0; background: #e9e9ee; }
.page { box-sizing: border-box; width: 21cm; max-width: calc(100% - 2rem);
  margin: 0 auto; padding: 3cm; background: #fff; color: #111;
  box-shadow: 0 1px 8px rgba(0,0,0,.15);
  font-family: "Times New Roman","BiauKai","標楷體","DFKai-SB","KaiTi","Noto Serif CJK TC",serif;
  font-size: 14pt; line-height: 1.8; text-align: justify; }
.cover { text-align: center; min-height: 24cm; display: flex; flex-direction: column;
  align-items: center; }
.cover > div { width: 100%; }
.cover .c-big { font-size: 20pt; margin: .15rem 0; }
.cover .c-title { font-size: 17pt; margin: .4rem 0; }
.cover .c-mid { font-size: 14pt; margin: .1rem 0; }
.cover .c-en { font-size: 12pt; margin: .1rem 0; }
.cover .c-gap { height: 2.2rem; }
.cover .c-fill { flex: 1; }
h2.chap { text-align: center; font-size: 18pt; margin: 2.4rem 0 1.6rem; font-weight: bold; }
h3 { font-size: 15pt; margin: 1.4rem 0 .6rem; }
h4 { font-size: 14pt; margin: 1.1rem 0 .5rem; }
p { margin: .5rem 0; text-indent: 2em; }
figure { text-align: center; margin: 1.2rem 0; }
figure img { max-width: 90%; }
figcaption { margin-top: .4rem; }
.tw-cap { text-align: center; margin: .6rem 0 .3rem; }
table { border-collapse: collapse; width: 100%; margin: .3rem 0 1rem; }
th, td { border: 1px solid #333; padding: .3rem .6rem; text-indent: 0; }
blockquote { border-left: 3px solid #ccc; margin: 1rem 0; padding-left: 1rem; color: #444; }
pre { background: #f5f5f5; padding: .8rem; overflow-x: auto; font-family: ui-monospace,monospace; text-indent: 0; }
.refs-title { page-break-before: always; }
.refs-group { font-weight: bold; margin: 1rem 0 .4rem; text-indent: 0; }
.ref { padding-left: 2em; text-indent: -2em; margin: .35rem 0; }
"""


def _html_cover(cover: dict, fallback_title: str) -> str:
    """Bilingual Taiwan-thesis cover block (mirrors _tw_cover_latex / _docx_cover)."""

    def g(k: str) -> str:
        return (cover.get(k) or "").strip()

    rows: list[str] = []

    def L(text: str, cls: str) -> None:
        rows.append(f'<div class="{cls}">{_h(text)}</div>')

    head_zh = " ".join(x for x in (g("universityZh"), g("departmentZh")) if x)
    if head_zh:
        L(head_zh, "c-big")
    if g("degreeZh"):
        L(g("degreeZh"), "c-big")
    for x in (g("departmentEn"), g("universityEn"), g("degreeEn")):
        if x:
            L(x, "c-en")
    L("", "c-gap")
    L(g("titleZh") or fallback_title, "c-title")
    if g("titleEn"):
        L(g("titleEn"), "c-en")
    L("", "c-gap")
    if g("advisorZh"):
        L(f"指導教授：{g('advisorZh')}", "c-mid")
    if g("advisorEn"):
        L(f"Advisor: {g('advisorEn')}", "c-en")
    if g("studentZh"):
        L(f"研究生：{g('studentZh')}", "c-mid")
    if g("studentEn"):
        L(f"Graduate Student: {g('studentEn')}", "c-en")
    rows.append('<div class="c-fill"></div>')
    if g("dateZh"):
        L(g("dateZh"), "c-mid")
    if g("dateEn"):
        L(g("dateEn"), "c-en")
    return '<section class="cover">' + "".join(rows) + "</section>"


def _html_refs_grouped(resolved: list[dict], refs_label: str, locale: str) -> str:
    """Grouped (網路資源/中文/英文) hanging-indent reference list as HTML."""
    labels = _REF_GROUP_LABELS.get(locale, _REF_GROUP_LABELS["en"])
    groups: dict[str, list[dict]] = {"web": [], "zh": [], "en": []}
    for c in resolved:
        groups[_ref_group(c)].append(c)
    out = [f'<h2 class="chap refs-title">{_h(refs_label)}</h2>']
    for key in ("web", "zh", "en"):
        g = sorted(groups[key], key=_ref_sort_key)
        if not g:
            continue
        out.append(f'<div class="refs-group">{_h(labels[key])}</div>')
        for c in g:
            out.append(f'<div class="ref">{_h(_tw_reference(c))}</div>')
    return "".join(out)


def to_html(
    document: dict,
    style: str,
    refs_label: str,
    template: str = "article",
    locale: str = "zh-Hant",
    cover: dict | None = None,
) -> str:
    content = _prepared_content(document)
    title = document.get("title") or "(untitled)"
    citations = collect_citations(content)
    order = [_cite_key(c) for c in citations]
    tw = template == "twthesis"
    cover_active = bool(tw and cover and any((v or "").strip() for v in cover.values()))
    ctx = {"chapter": 0, "fig": 0, "tab": 0}

    parts: list[str] = []
    if cover_active:
        parts.append(_html_cover(cover, title))
    elif tw:
        parts.append(f'<h2 class="chap">{_h(title)}</h2>')
    else:
        parts.append(f"<h1>{_h(title)}</h1>")

    refs_inline = False
    for b in _children(content):
        t = b.get("type")
        if t == "heading":
            lvl = (b.get("attrs") or {}).get("level", 1)
            if tw and lvl == 1:
                inner = _inline_html(_children(b), style, order)
                if _is_unnumbered_chapter(_plain_text(b)):
                    parts.append(f'<h2 class="chap">{inner}</h2>')
                else:
                    ctx["chapter"] += 1
                    ctx["fig"] = 0
                    ctx["tab"] = 0
                    parts.append(
                        f'<h2 class="chap">第{_cjk_num(ctx["chapter"])}章　{inner}</h2>'
                    )
            else:
                hl = min(lvl + 1, 4)
                parts.append(
                    f"<h{hl}>{_inline_html(_children(b), style, order)}</h{hl}>"
                )
        elif t in ("bulletList", "orderedList"):
            tag = "ul" if t == "bulletList" else "ol"
            items = "".join(
                "<li>"
                + "".join(
                    _inline_html(_children(ch), style, order) for ch in _children(item)
                )
                + "</li>"
                for item in _children(b)
            )
            parts.append(f"<{tag}>{items}</{tag}>")
        elif t == "blockquote":
            inner = "".join(
                "<p>" + _inline_html(_children(ch), style, order) + "</p>"
                for ch in _children(b)
            )
            parts.append(f"<blockquote>{inner}</blockquote>")
        elif t == "codeBlock":
            parts.append(f"<pre><code>{_h(_block_text(b))}</code></pre>")
        elif t == "horizontalRule":
            parts.append("<hr>")
        elif t == "figure":
            a = b.get("attrs") or {}
            cap = a.get("caption", "")
            if cap and tw:
                cap = f"圖 {_docx_caption_num(ctx, 'fig')}　{cap}"
            caphtml = f"<figcaption>{_h(cap)}</figcaption>" if cap else ""
            uri = _img_data_uri(a.get("src", ""))
            parts.append(
                f'<figure><img src="{_h(uri)}" alt="{_h(a.get("caption", ""))}">{caphtml}</figure>'
            )
        elif t == "mathBlock":
            parts.append(
                "<p>\\[" + _h((b.get("attrs") or {}).get("latex", "")) + "\\]</p>"
            )
        elif t == "tableBlock":
            cap, rows = _table_parts(b)
            tcap = ""
            if cap:
                label = f"表 {_docx_caption_num(ctx, 'tab')}　{cap}" if tw else cap
                tcap = (
                    f'<div class="tw-cap">{_h(label)}</div>'
                    if tw
                    else f"<caption>{_h(cap)}</caption>"
                )
            body = "".join(
                "<tr>" + "".join(f"<td>{_h(c)}</td>" for c in r) + "</tr>" for r in rows
            )
            if tw:
                parts.append(f"{tcap}<table>{body}</table>")
            else:
                parts.append(f"<table>{tcap}{body}</table>")
        elif t == "referenceList":
            resolved = [c for c in citations if not c.get("unlinked")]
            if resolved:
                refs_inline = True
                if tw:
                    parts.append(_html_refs_grouped(resolved, refs_label, locale))
                else:
                    items = "".join(
                        f'<li value="{_citation_number(order, _cite_key(c))}">'
                        f"{_h(full_reference(c, style, _citation_number(order, _cite_key(c))))}</li>"
                        for c in resolved
                    )
                    parts.append(
                        f'<h2>{_h(refs_label)}</h2><ol class="refs">{items}</ol>'
                    )
        elif t == "tableList":
            continue
        else:
            parts.append(f"<p>{_inline_html(_children(b), style, order)}</p>")

    resolved = [c for c in citations if not c.get("unlinked")]
    if resolved and not refs_inline:
        if tw:
            parts.append(_html_refs_grouped(resolved, refs_label, locale))
        else:
            items = "".join(
                f'<li value="{_citation_number(order, _cite_key(c))}">'
                f"{_h(full_reference(c, style, _citation_number(order, _cite_key(c))))}</li>"
                for c in resolved
            )
            parts.append(f'<h2>{_h(refs_label)}</h2><ol class="refs">{items}</ol>')

    body_html = "".join(parts)
    if tw:
        body_html = f'<div class="page">{body_html}</div>'
    css = _HTML_CSS_TW if tw else _HTML_CSS
    mathjax = (
        '<script>window.MathJax={tex:{inlineMath:[["\\\\(","\\\\)"]],displayMath:[["\\\\[","\\\\]"]]}};</script>'
        '<script async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>'
    )
    return (
        '<!doctype html><html lang="zh-Hant"><head><meta charset="utf-8">'
        f"<title>{_h(title)}</title><style>{css}</style>{mathjax}</head>"
        f"<body>{body_html}</body></html>"
    )
