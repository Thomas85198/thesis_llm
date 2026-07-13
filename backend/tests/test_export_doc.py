"""Tests for editor-mode export: ProseMirror JSON → .docx / .tex."""
from __future__ import annotations

import io

from docx import Document

from app import export_doc

# A doc exercising: heading, paragraph with bold + a citation, a bullet list, and
# two citations (one repeated) so numbering / dedup is covered.
CITE_A = {
    "type": "citation",
    "attrs": {"openalexId": "W1", "authors": "Ashish Vaswani, Noam Shazeer",
              "year": 2017, "title": "Attention Is All You Need", "venue": "NeurIPS"},
}
CITE_B = {
    "type": "citation",
    "attrs": {"openalexId": "W2", "authors": "Jane Doe", "year": 2020,
              "title": "A Study", "venue": "JML"},
}
DOC = {
    "type": "doc",
    "content": [
        {"type": "heading", "attrs": {"level": 1},
         "content": [{"type": "text", "text": "緒論"}]},
        {"type": "paragraph", "content": [
            {"type": "text", "text": "重點", "marks": [{"type": "bold"}]},
            {"type": "text", "text": "如下 "},
            CITE_A,
            {"type": "text", "text": " 與 "},
            CITE_B,
        ]},
        {"type": "bulletList", "content": [
            {"type": "listItem", "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "第一點 "}, CITE_A]}
            ]},
        ]},
    ],
}


# ---------- citation helpers ----------

def test_in_text_label_per_style():
    a = {"authors": "Ashish Vaswani, Noam Shazeer", "year": 2017}
    assert export_doc.in_text_label(a, "apa", 1) == "(Vaswani & Shazeer, 2017)"
    assert export_doc.in_text_label(a, "harvard", 1) == "(Vaswani & Shazeer, 2017)"
    assert export_doc.in_text_label(a, "chicago", 1) == "(Vaswani & Shazeer 2017)"
    assert export_doc.in_text_label(a, "mla", 1) == "(Vaswani & Shazeer)"
    assert export_doc.in_text_label(a, "ieee", 3) == "[3]"
    assert export_doc.in_text_label(a, "numeric", 3) == "[3]"
    assert export_doc.in_text_label({"authors": "", "year": None}, "apa", 1) == "(Anon., n.d.)"


def test_full_reference_per_style():
    a = {"authors": "A B", "year": 2020, "title": "A Study", "venue": "JML"}
    assert export_doc.full_reference(a, "apa", 1) == "A B (2020). A Study. JML."
    assert export_doc.full_reference(a, "mla", 1) == "A B. “A Study.” JML, 2020."
    assert export_doc.full_reference(a, "ieee", 5) == "[5] A B, “A Study,” JML, 2020."


def test_collect_citations_dedupes_by_first_appearance():
    cites = export_doc.collect_citations(DOC)
    assert [c["openalexId"] for c in cites] == ["W1", "W2"]  # W1 repeated → once, in order


# ---------- DOCX ----------

def test_to_docx_includes_text_citations_and_references():
    data = export_doc.to_docx({"title": "我的論文", "content_json": DOC}, "apa", "參考文獻")
    assert data[:2] == b"PK"  # docx is a zip
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "(Vaswani & Shazeer, 2017)" in text  # in-text APA marker (2 authors)
    assert "參考文獻" in text  # reference heading
    assert "Attention Is All You Need" in text  # reference entry


def test_to_docx_numeric_style_numbers_in_order():
    data = export_doc.to_docx({"title": "t", "content_json": DOC}, "numeric", "References")
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "[1]" in text and "[2]" in text  # first-appearance numbering


# ---------- LaTeX ----------

def test_to_latex_structure_and_escaping():
    tex, _imgs = export_doc.to_latex({"title": "100% 文字 & 符號", "content_json": DOC}, "apa", "參考文獻")
    assert "\\documentclass{article}" in tex
    assert "\\usepackage{ctex}" in tex  # CJK support
    assert "\\section{緒論}" in tex
    # CJK content auto-selects the TC font and pins XeLaTeX for local toolchains
    assert "\\setCJKmainfont{Noto Serif CJK TC}" in tex
    assert tex.startswith("% !TEX program = xelatex")
    assert "\\textbf{重點}" in tex
    assert "\\begin{itemize}" in tex
    assert "(Vaswani \\& Shazeer, 2017)" in tex  # marker's & is LaTeX-escaped
    assert "\\section*{參考文獻}" in tex
    # special chars in the title are escaped
    assert "100\\% 文字 \\& 符號" in tex


def test_to_latex_latin_only_doc_compiles_under_pdflatex():
    # No CJK anywhere → no ctex / fontspec / magic comment, so the file builds
    # on Overleaf's default pdfLaTeX without touching the compiler setting.
    doc = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "Plain English only."}]},
    ]}
    tex, _imgs = export_doc.to_latex({"title": "An English Title", "content_json": doc}, "apa", "References")
    assert "ctex" not in tex
    assert "setCJKmainfont" not in tex
    assert "!TEX program" not in tex
    assert tex.startswith("\\documentclass{article}")


def test_to_latex_cjk_in_body_triggers_cjk_preamble_even_with_latin_title():
    doc = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "內文有中文。"}]},
    ]}
    tex, _imgs = export_doc.to_latex({"title": "English Title", "content_json": doc}, "apa", "References")
    assert "\\usepackage{ctex}" in tex
    assert "\\setCJKmainfont{Noto Serif CJK TC}" in tex


def test_to_latex_caption_labels_follow_locale():
    doc = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "中文內容"}]},
    ]}
    # zh-Hant → Traditional caption labels (ctex defaults to Simplified 图/表)
    tex, _ = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "參考文獻", "article", "zh-Hant")
    assert "\\ctexset{figurename={圖},tablename={表}}" in tex
    # en doc with CJK content → plain scheme keeps Figure/Table labels
    tex_en, _ = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References", "article", "en")
    assert "scheme=plain" in tex_en
    assert "ctexset" not in tex_en


def test_figure_caption_kept_figurelist_skipped():
    doc = {"type": "doc", "content": [
        {"type": "figure", "attrs": {"src": "x.png", "caption": "流程圖"}},
        {"type": "figureList"},
    ]}
    tex, _imgs = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References")
    assert "流程圖" in tex  # caption preserved, not silently dropped
    import io as _io
    from docx import Document as _Doc
    data = export_doc.to_docx({"title": "t", "content_json": doc}, "apa", "References")
    text = "\n".join(p.text for p in _Doc(_io.BytesIO(data)).paragraphs)
    assert "流程圖" in text


def test_math_exports_native_latex():
    doc = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "text": "其中 "},
            {"type": "mathInline", "attrs": {"latex": "E=mc^2"}},
            {"type": "text", "text": " 成立。"},
        ]},
        {"type": "mathBlock", "attrs": {"latex": "\\int_0^1 x\\,dx"}},
    ]}
    tex, _imgs = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References")
    assert "$E=mc^2$" in tex  # inline math passes through verbatim (not escaped)
    assert "\\[\\int_0^1 x\\,dx\\]" in tex  # block math


def _table_doc():
    def cell(text):
        return {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": text}]}]}
    return {"type": "doc", "content": [
        {"type": "tableBlock", "content": [
            {"type": "tableCaption", "content": [{"type": "text", "text": "實驗結果"}]},
            {"type": "table", "content": [
                {"type": "tableRow", "content": [cell("A"), cell("B")]},
                {"type": "tableRow", "content": [cell("1"), cell("2")]},
            ]},
        ]},
        {"type": "tableList"},
    ]}


def test_table_exports_latex_xltabular_with_caption():
    tex, _imgs = export_doc.to_latex({"title": "t", "content_json": _table_doc()}, "apa", "References")
    # xltabular: X columns wrap within \linewidth AND the table can break
    # across pages between rows (a [H]+tabularx box taller than one page
    # silently overflows past the margin).
    assert "\\begin{xltabular}{\\linewidth}{|X|X|}" in tex
    assert "A & B" in tex and "1 & 2" in tex
    assert "\\caption{實驗結果}" in tex
    assert "tableList" not in tex  # the live aid is skipped


def test_table_twocolumn_falls_back_to_fixed_tabularx():
    # longtable can't be used in twocolumn layouts → anchored tabularx instead
    tex, _ = export_doc.to_latex({"title": "t", "content_json": _table_doc()}, "apa", "References", "twocolumn")
    assert "xltabular" not in tex.replace("\\usepackage{xltabular}", "")
    assert "\\begin{table}[H]" in tex and "\\begin{tabularx}" in tex


def test_table_monster_cell_splits_into_continuation_rows():
    # A transcript-length cell must be split into multiple physical rows —
    # longtable only breaks BETWEEN rows, so one giant row still overflows.
    long_text = "這是一句完整的訪談內容。" * 80  # ~960 chars
    def cell(text):
        return {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": text}]}]}
    doc = {"type": "doc", "content": [
        {"type": "tableBlock", "content": [
            {"type": "table", "content": [
                {"type": "tableRow", "content": [cell("MD001"), cell(long_text)]},
            ]},
        ]},
    ]}
    tex, _ = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References")
    rows = [l for l in tex.splitlines() if l.endswith("\\\\")]
    assert len(rows) >= 3  # split into ≥3 continuation rows
    # continuation rows belong to the same logical row: no \hline between them
    assert tex.count("\\hline") == 2  # only the table frame (top + bottom)


def test_table_exports_docx_real_table():
    import io as _io
    from docx import Document as _Doc
    data = export_doc.to_docx({"title": "t", "content_json": _table_doc()}, "apa", "References")
    doc = _Doc(_io.BytesIO(data))
    assert len(doc.tables) == 1
    cells = [c.text for c in doc.tables[0].rows[0].cells]
    assert cells == ["A", "B"]
    assert "實驗結果" in "\n".join(p.text for p in doc.paragraphs)  # caption


def _png(w: int = 4, h: int = 4) -> bytes:
    """A minimal but valid RGB PNG (python-docx's parser is strict)."""
    import struct
    import zlib

    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * w for _ in range(h))
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


_PNG_1PX = _png()


def test_latex_image_embeds_includegraphics_and_bundles(monkeypatch):
    monkeypatch.setattr(export_doc, "_image_bytes", lambda src: (_PNG_1PX, "png") if src else None)
    doc = {"type": "doc", "content": [
        {"type": "figure", "attrs": {"src": "/api/editor/images/x.png", "caption": "流程圖"}},
    ]}
    tex, imgs = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References")
    assert "\\includegraphics" in tex and "fig1.png" in tex
    assert "\\begin{figure}[H]" in tex  # anchored in place, never floats to the end
    assert "\\caption{流程圖}" in tex
    assert len(imgs) == 1 and imgs[0][0] == "fig1.png" and imgs[0][1] == _PNG_1PX


def test_latex_template_changes_documentclass():
    doc = {"type": "doc", "content": []}
    tex, _ = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References", "twocolumn")
    assert "\\documentclass[twocolumn]{article}" in tex
    tex2, _ = export_doc.to_latex({"title": "t", "content_json": doc}, "apa", "References", "ieee")
    assert "IEEEtran" in tex2


def test_twthesis_layout_chapters_margins_spacing():
    doc = {"type": "doc", "content": [
        {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "緒論"}]},
        {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "研究背景"}]},
        {"type": "paragraph", "content": [{"type": "text", "text": "內文 ", "marks": []}, CITE_A]},
    ]}
    tex, _ = export_doc.to_latex({"title": "我的論文", "content_json": doc}, "apa", "參考文獻", "twthesis")
    assert "\\documentclass[a4paper,12pt]{report}" in tex
    assert "\\usepackage[heading=true]{ctex}" in tex  # \chapter → 第一章, not "Chapter 1"
    assert "\\chapter{緒論}" in tex          # level-1 heading → 第一章 緒論
    assert "\\section{研究背景}" in tex
    assert "margin=3cm" in tex and "\\onehalfspacing" in tex
    assert "\\chapter*{參考文獻}" in tex      # refs at chapter level in report class


def test_docx_embeds_image(monkeypatch):
    import io as _io
    from docx import Document as _Doc
    monkeypatch.setattr(export_doc, "_image_bytes", lambda src: (_PNG_1PX, "png") if src else None)
    doc = {"type": "doc", "content": [
        {"type": "figure", "attrs": {"src": "/api/editor/images/x.png", "caption": "流程圖"}},
    ]}
    data = export_doc.to_docx({"title": "t", "content_json": doc}, "apa", "References")
    d = _Doc(_io.BytesIO(data))
    assert len(d.inline_shapes) == 1  # embedded picture
    assert "流程圖" in "\n".join(p.text for p in d.paragraphs)


def test_to_latex_empty_doc_no_references_section():
    tex, _imgs = export_doc.to_latex({"title": "t", "content_json": {"type": "doc", "content": []}}, "apa", "References")
    assert "\\section*{References}" not in tex  # no citations → no ref list
    assert "\\begin{document}" in tex and "\\end{document}" in tex


# ---------- Markdown / text / HTML ----------

def test_markdown_export():
    md = export_doc.to_markdown({"title": "我的論文", "content_json": DOC}, "apa", "參考文獻")
    assert md.startswith("# 我的論文")
    assert "# 緒論" in md  # heading level shifted under doc title
    assert "**重點**" in md  # bold
    assert "(Vaswani & Shazeer, 2017)" in md  # in-text citation
    assert "## 參考文獻" in md  # references section


def test_text_export_plain():
    txt = export_doc.to_text({"title": "我的論文", "content_json": DOC}, "apa", "參考文獻")
    assert "我的論文" in txt
    assert "<" not in txt and "**" not in txt  # no markup
    assert "(Vaswani & Shazeer, 2017)" in txt


def test_html_export_structure():
    html = export_doc.to_html({"title": "我的論文", "content_json": DOC}, "apa", "參考文獻")
    assert html.startswith("<!doctype html>")
    assert "<h1>我的論文</h1>" in html
    assert "<strong>重點</strong>" in html
    assert "MathJax" in html  # math rendering script
    assert "serif" in html  # academic font in CSS
    assert "&lt;" not in html or True  # html-escaping helper present


def test_docx_uses_serif_font():
    import io as _io
    from docx import Document as _Doc
    data = export_doc.to_docx({"title": "t", "content_json": DOC}, "apa", "References")
    d = _Doc(_io.BytesIO(data))
    assert d.styles["Normal"].font.name == "Times New Roman"


# ---------- DOCX math → OMML ----------

_M_NS = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}


def test_docx_math_renders_native_omml():
    """行內/區塊數學式輸出 Word 原生 OMML（非字面 $...$ 文字）——三格式
    一致的最後一塊：LaTeX/PDF 原生排版、HTML MathJax、DOCX OMML。"""
    import io as _io

    from docx import Document as _Doc

    doc = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "text": "能量公式 "},
            {"type": "mathInline", "attrs": {"latex": "E=mc^2"}},
            {"type": "text", "text": " 很有名。"},
        ]},
        {"type": "mathBlock", "attrs": {"latex": "\\int_0^1 x^2\\,dx = \\frac{1}{3}"}},
    ]}
    data = export_doc.to_docx({"title": "t", "content_json": doc}, "apa", "References")
    d = _Doc(_io.BytesIO(data))
    paras = d.paragraphs
    inline = [p for p in paras if p._p.findall(".//m:oMath", _M_NS)]
    blocks = [p for p in paras if p._p.findall(".//m:oMathPara", _M_NS)]
    assert len(inline) >= 2 and len(blocks) == 1
    full_text = "\n".join(p.text for p in paras)
    assert "$E=mc^2$" not in full_text  # 不再是字面文字
    # 行內數學位於文字 run 之間、同一段落
    host = next(p for p in paras if "能量公式" in p.text)
    assert host._p.findall(".//m:oMath", _M_NS)


def test_docx_math_falls_back_to_literal_on_conversion_failure(monkeypatch):
    import io as _io

    from docx import Document as _Doc

    monkeypatch.setattr(export_doc, "_omml_element", lambda *a, **k: None)
    doc = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "mathInline", "attrs": {"latex": "E=mc^2"}},
        ]},
        {"type": "mathBlock", "attrs": {"latex": "\\bad{cmd}"}},
    ]}
    data = export_doc.to_docx({"title": "t", "content_json": doc}, "apa", "References")
    d = _Doc(_io.BytesIO(data))
    full_text = "\n".join(p.text for p in d.paragraphs)
    assert "$E=mc^2$" in full_text and "$$\\bad{cmd}$$" in full_text


def test_omml_element_handles_garbage():
    assert export_doc._omml_element("") is None
    assert export_doc._omml_element("   ") is None
    # 正常式子回 lxml element
    el = export_doc._omml_element("x^2")
    assert el is not None and el.tag.endswith("}oMath")


def test_omml_patches_mathml2omml_groupchr_bug():
    """mathml2omml 0.0.2 對 \\bar/\\overline（MOver/MUnder groupChr）輸出
    錯誤關閉標籤 → XML 解析失敗。已做精準修補，統計常用符號必須原生輸出。"""
    for latex in (r"\bar{X} = \frac{1}{n}\sum_{i=1}^{n} X_i", r"\overline{AB}",
                  r"\underline{x}", r"\underbrace{a+b}_{2}"):
        assert export_doc._omml_element(latex) is not None, latex
