"""Tests for editor-mode import: .txt / .md / .docx / .tex → ProseMirror JSON.

The inverse of test_export_doc. Verifies each parser produces nodes whose names
and attrs match the editor schema, plus an export→import round-trip on markdown
(the lossless-ish path).
"""

from __future__ import annotations

import io

from docx import Document

from app import export_doc, import_doc


# ---------- helpers ----------


def _types(doc: dict) -> list[str]:
    return [b.get("type") for b in doc["content"]]


def _find(node: dict, type_name: str) -> list[dict]:
    """All descendant nodes of a given type (depth-first)."""
    out: list[dict] = []

    def walk(n: dict) -> None:
        if n.get("type") == type_name:
            out.append(n)
        for ch in n.get("content", []):
            walk(ch)

    walk(node)
    return out


def _all_text(node: dict) -> str:
    if node.get("type") == "text":
        return node.get("text", "")
    return "".join(_all_text(ch) for ch in node.get("content", []))


# ---------- plain text ----------


def test_text_paragraphs_and_hardbreaks():
    _, doc = import_doc.from_text("First para line one\nline two\n\nSecond para")
    assert _types(doc) == ["paragraph", "paragraph"]
    p0 = doc["content"][0]
    assert any(n["type"] == "hardBreak" for n in p0["content"])
    assert _all_text(p0) == "First para line oneline two"  # break between the two lines


def test_text_empty_gives_one_paragraph():
    _, doc = import_doc.from_text("   \n\n  ")
    assert doc["content"] == [{"type": "paragraph"}]


# ---------- markdown ----------


def test_markdown_title_extracted_and_dropped():
    title, doc = import_doc.from_markdown("# My Paper\n\nHello world.")
    assert title == "My Paper"
    assert _types(doc) == ["paragraph"]  # the H1 is consumed as the title


def test_markdown_headings_shift_up_when_h1_becomes_title():
    # "# title + ## chapters" convention: once the leading H1 is consumed as
    # the title, ## must become H1 (else the Taiwan-thesis export has no
    # chapters and numbers sections 0.1/0.2).
    title, doc = import_doc.from_markdown("# T\n\n## 緒論\n\ntext\n\n### 背景\n")
    assert title == "T"
    levels = [
        (b.get("attrs") or {}).get("level")
        for b in doc["content"]
        if b["type"] == "heading"
    ]
    assert levels == [1, 2]


def test_markdown_headings_not_shifted_without_title_h1():
    _, doc = import_doc.from_markdown("## 緒論\n\ntext\n")
    levels = [
        (b.get("attrs") or {}).get("level")
        for b in doc["content"]
        if b["type"] == "heading"
    ]
    assert levels == [2]  # no title consumed → levels stay as authored


def test_markdown_marks_and_inline_math():
    _, doc = import_doc.from_markdown("A **b** _i_ `c` ~~s~~ and $E=mc^2$ end.")
    para = doc["content"][0]
    marks = {
        tuple(sorted(m["type"] for m in n.get("marks", [])))
        for n in para["content"]
        if n["type"] == "text"
    }
    assert ("bold",) in marks and ("italic",) in marks
    assert ("code",) in marks and ("strike",) in marks
    assert any(
        n["type"] == "mathInline" and n["attrs"]["latex"] == "E=mc^2"
        for n in para["content"]
    )


def test_markdown_list_quote_code_hr_mathblock():
    src = "- a\n- b\n\n> quoted\n\n```\ncode\n```\n\n---\n\n$$\n\\int x\n$$"
    _, doc = import_doc.from_markdown(src)
    assert "bulletList" in _types(doc)
    assert "blockquote" in _types(doc)
    assert "codeBlock" in _types(doc)
    assert "horizontalRule" in _types(doc)
    mb = _find(doc, "mathBlock")
    assert mb and mb[0]["attrs"]["latex"] == "\\int x"
    li = _find(doc, "listItem")
    assert len(li) == 2


def test_markdown_ordered_list():
    _, doc = import_doc.from_markdown("1. one\n2. two\n3. three")
    assert _types(doc) == ["orderedList"]
    assert len(_find(doc, "listItem")) == 3


def test_markdown_image_becomes_figure():
    _, doc = import_doc.from_markdown("![a caption](http://x/y.png)")
    figs = _find(doc, "figure")
    assert len(figs) == 1
    assert figs[0]["attrs"]["src"] == "http://x/y.png"
    assert figs[0]["attrs"]["caption"] == "a caption"


def test_markdown_table_structure():
    _, doc = import_doc.from_markdown("| H1 | H2 |\n| -- | -- |\n| a | b |\n| c | d |")
    tb = _find(doc, "tableBlock")
    assert tb
    assert _find(doc, "tableCaption")  # present (possibly empty)
    headers = _find(doc, "tableHeader")
    cells = _find(doc, "tableCell")
    assert len(headers) == 2  # header row
    assert len(cells) == 4  # two body rows × two cols
    assert _all_text(headers[0]) == "H1"
    assert _all_text(cells[3]) == "d"


# ---------- docx ----------


def _make_docx() -> bytes:
    d = Document()
    d.add_heading("Doc Title", level=0)  # Title style
    d.add_heading("Section One", level=1)
    p = d.add_paragraph()
    p.add_run("normal ")
    p.add_run("bold").bold = True
    d.add_paragraph("first bullet", style="List Bullet")
    d.add_paragraph("second bullet", style="List Bullet")
    d.add_paragraph("a quote", style="Quote")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "H1"
    t.cell(0, 1).text = "H2"
    t.cell(1, 0).text = "a"
    t.cell(1, 1).text = "b"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_structure_and_title():
    title, doc = import_doc.from_docx(_make_docx())
    assert title == "Doc Title"
    types = _types(doc)
    assert "heading" in types
    assert "bulletList" in types
    assert "blockquote" in types
    assert "tableBlock" in types


def test_docx_bold_run_and_heading_level():
    _, doc = import_doc.from_docx(_make_docx())
    h = _find(doc, "heading")[0]
    assert h["attrs"]["level"] == 1
    assert _all_text(h) == "Section One"
    # the bold run keeps a bold mark
    bolds = [
        n
        for n in _find(doc, "text")
        if any(m["type"] == "bold" for m in n.get("marks", []))
    ]
    assert any(n["text"] == "bold" for n in bolds)


def test_docx_list_grouping():
    _, doc = import_doc.from_docx(_make_docx())
    lists = _find(doc, "bulletList")
    assert len(lists) == 1  # two items collapse into ONE list
    assert len(_find(doc, "listItem")) == 2


# ---------- latex ----------


def test_latex_sections_and_marks():
    src = (
        r"\documentclass{article}\title{LaTeX Paper}\begin{document}\maketitle"
        "\n\\section{Intro}\nSome \\textbf{bold} and \\textit{italic} text.\n"
        "\\subsection{Sub}\nMore.\n\\end{document}"
    )
    title, doc = import_doc.from_latex(src)
    assert title == "LaTeX Paper"
    hs = _find(doc, "heading")
    assert hs[0]["attrs"]["level"] == 1 and _all_text(hs[0]) == "Intro"
    assert hs[1]["attrs"]["level"] == 2
    bolds = [
        n
        for n in _find(doc, "text")
        if any(m["type"] == "bold" for m in n.get("marks", []))
    ]
    assert any(n["text"] == "bold" for n in bolds)


def test_latex_itemize_and_inline_math():
    src = "\\begin{itemize}\n\\item first\n\\item second\n\\end{itemize}\nText $a^2+b^2$ here."
    _, doc = import_doc.from_latex(src)
    assert "bulletList" in _types(doc)
    assert len(_find(doc, "listItem")) == 2
    mi = _find(doc, "mathInline")
    assert mi and mi[0]["attrs"]["latex"] == "a^2+b^2"


def test_latex_table_and_figure():
    src = (
        r"\begin{table}\caption{My Table}"
        "\n\\begin{tabular}{|l|l|}\n\\hline\nA & B \\\\\n\\hline\nc & d \\\\\n\\hline\n"
        "\\end{tabular}\n\\end{table}\n"
        r"\begin{figure}\includegraphics{plot.png}\caption{My Figure}\end{figure}"
    )
    _, doc = import_doc.from_latex(src)
    tb = _find(doc, "tableBlock")
    assert tb
    assert _all_text(_find(doc, "tableCaption")[0]) == "My Table"
    assert len(_find(doc, "tableHeader")) == 2
    fig = _find(doc, "figure")
    assert fig and fig[0]["attrs"]["caption"] == "My Figure"


def test_latex_escapes():
    _, doc = import_doc.from_latex(r"100\% of a\_b \& c.")
    assert _all_text(doc) == "100% of a_b & c."


# ---------- dispatch ----------


def test_dispatch_by_extension_and_title_fallback():
    title, doc = import_doc.to_prosemirror("notes.txt", b"hello")
    assert title == "notes"  # filename stem fallback
    assert doc["type"] == "doc"
    title2, _ = import_doc.to_prosemirror("paper.md", b"# Real Title\n\nbody")
    assert title2 == "Real Title"


def test_dispatch_rejects_unknown():
    try:
        import_doc.to_prosemirror("x.xyz", b"whatever")
        assert False, "should have raised"
    except ValueError:
        pass


# ---------- round-trip: export → import preserves structure ----------


def test_markdown_roundtrip_preserves_blocks():
    original = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "緒論"}],
            },
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "重點", "marks": [{"type": "bold"}]},
                    {"type": "text", "text": "如下。"},
                ],
            },
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "一"}],
                            }
                        ],
                    },
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "二"}],
                            }
                        ],
                    },
                ],
            },
        ],
    }
    md = export_doc.to_markdown(
        {"title": "我的論文", "content_json": original}, "apa", "參考文獻"
    )
    title, doc = import_doc.from_markdown(md)
    assert title == "我的論文"
    assert "heading" in _types(doc)
    assert "bulletList" in _types(doc)
    assert len(_find(doc, "listItem")) == 2
    bolds = [
        n
        for n in _find(doc, "text")
        if any(m["type"] == "bold" for m in n.get("marks", []))
    ]
    assert any(n["text"] == "重點" for n in bolds)


# ---------- post-processing: heading demotion + directory re-linking ----------


def test_demote_pseudo_headings():
    blocks = [
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "緒論"}],
        },
        {
            "type": "heading",
            "attrs": {"level": 3},
            "content": [
                {
                    "type": "text",
                    "text": "這是一段很長的內文，說明研究背景與動機，並以句號結尾。",
                }
            ],
        },
        {
            "type": "heading",
            "attrs": {"level": 2},
            "content": [{"type": "text", "text": "研究方法。"}],
        },
    ]
    out = import_doc._demote_pseudo_headings(blocks)
    assert out[0]["type"] == "heading"  # short heading kept
    assert out[1]["type"] == "paragraph"  # long body prose demoted
    assert out[2]["type"] == "paragraph"  # ends with 。 → demoted


def test_relink_directories():
    blocks = [
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "目錄"}],
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "第一章\t"}],
        },  # tab debris
        {"type": "paragraph"},  # empty debris
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "圖目錄"}],
        },
        {"type": "paragraph"},
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "表目錄"}],
        },
        {"type": "paragraph"},
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "第一章 緒論"}],
        },
        {"type": "paragraph", "content": [{"type": "text", "text": "正文開始"}]},
    ]
    out = import_doc._relink_directories(blocks)
    assert [b["type"] for b in out] == [
        "heading",
        "tableOfContents",  # 目錄 → live TOC, debris dropped
        "heading",
        "figureList",  # 圖目錄 (matched before 目錄)
        "heading",
        "tableList",  # 表目錄
        "heading",
        "paragraph",  # real content untouched
    ]


def test_relink_keeps_leader_dot_page_numbers_out():
    blocks = [
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "目錄"}],
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "第一章 緒論……12"}],
        },  # leader dots + page
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "這是真正的正文段落"}],
        },
    ]
    out = import_doc._relink_directories(blocks)
    assert [b["type"] for b in out] == ["heading", "tableOfContents", "paragraph"]
    assert _all_text(out[2]) == "這是真正的正文段落"


def test_unwrap_pseudo_blockquotes():
    blocks = [
        {
            "type": "blockquote",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {
                            "type": "text",
                            "text": "Bhattacherjee, A. (2001). Understanding IS continuance. MIS Quarterly, 25(3).",
                        }
                    ],
                }
            ],
        },
        {
            "type": "blockquote",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "知識就是力量"}],
                }
            ],
        },  # short pull-quote, no full stop
    ]
    out = import_doc._unwrap_pseudo_blockquotes(blocks)
    assert out[0]["type"] == "paragraph"  # reference entry unwrapped
    assert out[1]["type"] == "blockquote"  # genuine short quote kept


def test_attach_figure_and_table_captions():
    blocks = [
        {"type": "figure", "attrs": {"src": "/x.png", "caption": "", "alt": ""}},
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "圖1-1 2018全球社群媒體使用概況"}],
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "表2-1 社群媒體分類"}],
        },
        {
            "type": "tableBlock",
            "content": [{"type": "tableCaption"}, {"type": "table", "content": []}],
        },
    ]
    out = import_doc._attach_captions(blocks)
    assert [b["type"] for b in out] == [
        "figure",
        "tableBlock",
    ]  # caption paragraphs folded in
    assert (
        out[0]["attrs"]["caption"] == "2018全球社群媒體使用概況"
    )  # 圖 prefix stripped
    assert _all_text(out[1]["content"][0]) == "社群媒體分類"  # tableCaption set


def test_attach_caption_styled_as_heading():
    # theses often style the figure caption as a heading — must still be folded in
    blocks = [
        {"type": "figure", "attrs": {"src": "/x.png", "caption": "", "alt": ""}},
        {
            "type": "heading",
            "attrs": {"level": 3},
            "content": [{"type": "text", "text": "圖3-1研究流程"}],
        },
    ]
    out = import_doc._attach_captions(blocks)
    assert len(out) == 1 and out[0]["type"] == "figure"
    assert out[0]["attrs"]["caption"] == "研究流程"


def test_attach_captions_ignores_non_caption_paragraph():
    blocks = [
        {"type": "figure", "attrs": {"src": "/x.png", "caption": "", "alt": ""}},
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "這是正文不是圖說。"}],
        },
    ]
    out = import_doc._attach_captions(blocks)
    assert (
        len(out) == 2 and out[0]["attrs"]["caption"] == ""
    )  # body paragraph left alone


def test_export_expands_table_of_contents():
    doc = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "目錄"}],
            },
            {"type": "tableOfContents"},
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "緒論"}],
            },
            {
                "type": "heading",
                "attrs": {"level": 2},
                "content": [{"type": "text", "text": "研究背景"}],
            },
        ],
    }
    md = export_doc.to_markdown({"title": "T", "content_json": doc}, "apa", "參考文獻")
    # the live node expands into a static list, so each heading appears twice
    assert md.count("緒論") >= 2 and md.count("研究背景") >= 2
